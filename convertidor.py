"""
╔══════════════════════════════════════════════════════════════════╗
║  SIAA — Convertidor Completo (Word + PDF + Excel)               ║
║  Seccional Bucaramanga — Rama Judicial                           ║
╠══════════════════════════════════════════════════════════════════╣
║  PROPÓSITO:                                                       ║
║    - .doc / .docx  → Markdown  (vía python-docx)                 ║
║    - .pdf          → Markdown  (vía LibreOffice headless         ║
║                                 o pymupdf4llm como fallback)     ║
║    - .xls / .xlsx  → tabla SQLite  (vía pandas)                  ║
║                                                                   ║
║  DIFERENCIA CON convertidor_pdf.py:                              ║
║    Este script procesa carpetas con mezcla de archivos.          ║
║    Cada carpeta = un "instructivo" con su Word y su Excel.       ║
║    convertidor_pdf.py solo convierte PDFs sueltos.               ║
║                                                                   ║
║  CAMBIOS vs versión Windows:                                      ║
║    ✗ PowerShell/Word COM       → ✓ LibreOffice headless          ║
║    ✗ C:\SIAA\...               → ✓ /opt/siaa/...                 ║
║    ✗ temp dir Windows          → ✓ /tmp/siaa_temp/               ║
║                                                                   ║
║  Instalación (una sola vez):                                      ║
║    sudo dnf install libreoffice-headless                          ║
║    pip install python-docx pandas openpyxl xlrd pymupdf4llm \\   ║
║                requests --break-system-packages                   ║
║                                                                   ║
║  Uso:                                                             ║
║    python3 convertidor.py                                         ║
║    python3 convertidor.py --origen /ruta --dest-md /salida       ║
║    python3 convertidor.py --only-folder "Juzgado Civil Municipal" ║
╚══════════════════════════════════════════════════════════════════╝
"""

import argparse
import os
import re
import shutil
import sqlite3
import subprocess
import tempfile
import time
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import pandas as pd

try:
    import docx
except ImportError:
    docx = None

# ================================================================
#  RUTAS POR DEFECTO — Linux / Fedora 43
#  Anteriores (Windows):
#    C:\SIAA\Instructivos     → /opt/siaa/instructivos
#    C:\SIAA\fuentes          → /opt/siaa/fuentes
#    C:\SIAA\institucional.db → /opt/siaa/institucional.db
#    C:\SIAA\conversion...log → /opt/siaa/logs/conversion_errores.log
# ================================================================

DEFAULT_ORIGEN  = Path("/opt/siaa/instructivos")            # Carpetas con Word + Excel
DEFAULT_DEST_MD = Path("/opt/siaa/fuentes")                 # Salida .md (colección general)
DEFAULT_DB      = Path("/opt/siaa/institucional.db")        # Base de datos SQLite
DEFAULT_LOG     = Path("/opt/siaa/logs/conversion_errores.log")
TEMP_DIR        = Path("/tmp/siaa_temp")                    # Antes: C:\Users\...\AppData\Temp

# Proxy SIAA — recarga el índice automáticamente al finalizar
PROXY_RECARGAR = "http://localhost:5000/siaa/recargar"
RECARGAR_AUTO  = True

# ================================================================
#  ESTRUCTURAS DE DATOS
# ================================================================

@dataclass
class FolderResult:
    folder_name: str
    slug:        str
    md_path:     Path
    md_ok:       bool       = False
    md_msg:      str        = ""
    sql_ok:      bool       = False
    sql_rows:    int        = 0
    sql_msg:     str        = ""
    errors:      list[str]  = field(default_factory=list)


# ================================================================
#  UTILIDADES DE TEXTO
# ================================================================

def slugify_ascii(name: str) -> str:
    normalized = unicodedata.normalize("NFKD", name)
    ascii_only = normalized.encode("ascii", "ignore").decode("ascii")
    collapsed  = re.sub(r"[\s\-/]+", "_", ascii_only.lower())
    cleaned    = re.sub(r"[^a-z0-9_]", "", collapsed)
    cleaned    = re.sub(r"_+", "_", cleaned).strip("_")
    return cleaned or "sin_nombre"


def sanitize_column(name: Any) -> str:
    return slugify_ascii("" if name is None else str(name)) or "columna"


def make_unique_columns(columns: list[str]) -> list[str]:
    used:   dict[str, int] = {}
    unique: list[str]      = []
    for col in columns:
        base  = col or "columna"
        count = used.get(base, 0) + 1
        used[base] = count
        unique.append(base if count == 1 else f"{base}_{count}")
    return unique


def markdown_table(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    width      = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header, body = normalized[0], normalized[1:]
    lines = [
        f"| {' | '.join(header)} |",
        f"| {' | '.join(['---'] * width)} |",
    ]
    lines.extend(f"| {' | '.join(r)} |" for r in body)
    return "\n".join(lines)


def _safe_cell(value: Any) -> str:
    if value is None:
        return ""
    return str(value).replace("\n", " ").strip().replace("|", "\\|")


# ================================================================
#  CONVERSIÓN WORD (.docx) → MARKDOWN
#  Sin cambios respecto a la versión Windows — python-docx es
#  multiplataforma y no depende de COM ni de Office instalado.
# ================================================================

def docx_to_markdown(docx_path: Path, folder_name: str) -> tuple[bool, str]:
    """Extrae texto y tablas de .docx a Markdown usando python-docx."""
    if docx is None:
        return False, "python-docx no instalado: pip install python-docx --break-system-packages"

    document = docx.Document(str(docx_path))
    lines: list[str] = [f"# {folder_name}", ""]

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (paragraph.style.name or "").lower()
        if "heading" in style_name:
            level_match = re.search(r"(\d+)", style_name)
            level = int(level_match.group(1)) if level_match else 2
            level = max(1, min(level, 6))
            lines.extend([f"{'#' * level} {text}", ""])
        else:
            lines.extend([text, ""])

    for table in document.tables:
        table_rows = [[_safe_cell(cell.text) for cell in row.cells] for row in table.rows]
        md_table = markdown_table(table_rows)
        if md_table:
            lines.extend([md_table, ""])

    if len(lines) <= 2:
        lines.extend([
            "## Estado de conversion", "",
            f"- Archivo Word detectado: `{docx_path.name}`",
            "- Resultado: OK sin contenido textual visible.", "",
        ])

    return True, "\n".join(lines).strip() + "\n"


# ================================================================
#  [LINUX FIX] CONVERSIÓN .doc / .pdf → .docx
#
#  ANTES (Windows): PowerShell + Word COM (New-Object Word.Application)
#  AHORA (Linux):   LibreOffice headless
#
#  LibreOffice convierte .doc y .pdf a .docx de forma nativa sin
#  necesidad de Microsoft Office instalado.
#
#  Instalación en Fedora:
#    sudo dnf install libreoffice-headless
#
#  Si LibreOffice no está disponible, se intenta pymupdf4llm como
#  fallback directo para PDFs (sin pasar por .docx).
# ================================================================

def _libreoffice_disponible() -> bool:
    """Verifica si LibreOffice headless está instalado en el sistema."""
    return shutil.which("libreoffice") is not None or shutil.which("soffice") is not None


def _cmd_libreoffice() -> str:
    """Devuelve el comando correcto: 'libreoffice' o 'soffice'."""
    if shutil.which("libreoffice"):
        return "libreoffice"
    if shutil.which("soffice"):
        return "soffice"
    return "libreoffice"  # fallback que fallará con mensaje claro


def convert_to_docx_via_libreoffice(input_path: Path, output_dir: Path) -> tuple[bool, Path | None, str]:
    """
    Convierte .doc o .pdf a .docx usando LibreOffice headless.

    Reemplaza completamente la función convert_to_docx_via_word
    que usaba PowerShell + Word COM en Windows.

    Args:
        input_path: archivo .doc o .pdf de entrada
        output_dir: directorio donde LibreOffice dejará el .docx resultante

    Returns:
        (exito, ruta_docx_resultante, mensaje)
    """
    output_dir.mkdir(parents=True, exist_ok=True)

    if not _libreoffice_disponible():
        return False, None, (
            "LibreOffice no está instalado. "
            "Instale con: sudo dnf install libreoffice-headless"
        )

    cmd = [
        _cmd_libreoffice(),
        "--headless",
        "--norestore",
        "--convert-to", "docx",
        "--outdir", str(output_dir),
        str(input_path),
    ]

    try:
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=120   # 2 minutos máximo por archivo
        )
    except subprocess.TimeoutExpired:
        return False, None, "LibreOffice tardó más de 2 minutos — archivo puede estar corrupto."
    except FileNotFoundError:
        return False, None, "LibreOffice no encontrado. Instale: sudo dnf install libreoffice-headless"

    if result.returncode != 0:
        detalle = result.stderr.strip() or result.stdout.strip() or "error desconocido"
        return False, None, f"LibreOffice error (código {result.returncode}): {detalle[:200]}"

    # LibreOffice guarda el .docx con el mismo nombre de stem en output_dir
    docx_esperado = output_dir / (input_path.stem + ".docx")
    if not docx_esperado.exists():
        # A veces LibreOffice cambia ligeramente el nombre — buscar en el dir
        candidatos = list(output_dir.glob("*.docx"))
        if candidatos:
            docx_esperado = max(candidatos, key=lambda p: p.stat().st_mtime)
        else:
            return False, None, "LibreOffice reportó éxito pero no generó archivo .docx."

    return True, docx_esperado, "Conversión a .docx completada con LibreOffice."


def convert_pdf_directo(pdf_path: Path) -> tuple[bool, str]:
    """
    Fallback: convierte PDF directamente a Markdown con pymupdf4llm,
    sin pasar por .docx. Más fiel para PDFs con tablas complejas.
    """
    try:
        import pymupdf4llm
        md_text = pymupdf4llm.to_markdown(str(pdf_path))
        return True, md_text
    except ImportError:
        return False, "pymupdf4llm no instalado: pip install pymupdf4llm --break-system-packages"
    except Exception as e:
        return False, f"Error en pymupdf4llm: {e}"


# ================================================================
#  LÓGICA DE CONVERSIÓN COMPLETA POR ARCHIVO
# ================================================================

def convert_source_to_md(source_path: Path, md_path: Path,
                          folder_name: str, convert_doc: bool) -> tuple[bool, str]:
    """
    Maneja la conversión completa de cualquier formato a Markdown.

    Estrategia según extensión:
      .docx → python-docx (directo, sin intermediario)
      .doc  → LibreOffice → .docx → python-docx
      .pdf  → pymupdf4llm directo (mejor calidad para tablas)
              fallback: LibreOffice → .docx → python-docx
    """
    suffix = source_path.suffix.lower()

    if suffix not in (".doc", ".docx", ".pdf"):
        msg = f"Extensión no soportada: {source_path.suffix}"
        _write_error_md(md_path, folder_name, source_path.name, msg)
        return False, msg

    # ── .docx: conversión directa con python-docx ─────────────
    if suffix == ".docx":
        ok, md_or_err = docx_to_markdown(source_path, folder_name)
        if ok:
            md_path.write_text(md_or_err, encoding="utf-8")
            return True, "Word .docx convertido a Markdown."
        _write_error_md(md_path, folder_name, source_path.name, md_or_err)
        return False, md_or_err

    # ── .pdf: pymupdf4llm directo (preferido en Linux) ─────────
    if suffix == ".pdf":
        ok_directo, md_o_err = convert_pdf_directo(source_path)
        if ok_directo:
            encabezado = f"# {folder_name}\n\n"
            md_path.write_text(encabezado + md_o_err, encoding="utf-8")
            return True, "PDF convertido a Markdown con pymupdf4llm."

        # Fallback: LibreOffice → .docx → python-docx
        print(f"     pymupdf4llm falló ({md_o_err[:60]}), intentando LibreOffice...")
        temp_dir = TEMP_DIR / f"{slugify_ascii(folder_name)}_{os.getpid()}"
        ok_lo, docx_path, msg_lo = convert_to_docx_via_libreoffice(source_path, temp_dir)
        if not ok_lo:
            _write_error_md(md_path, folder_name, source_path.name, msg_lo)
            return False, msg_lo
        ok, md_or_err = docx_to_markdown(docx_path, folder_name)
        if ok:
            md_path.write_text(md_or_err, encoding="utf-8")
            return True, "PDF convertido vía LibreOffice → .docx → Markdown."
        _write_error_md(md_path, folder_name, source_path.name, md_or_err)
        return False, md_or_err

    # ── .doc: LibreOffice → .docx → python-docx ────────────────
    if suffix == ".doc":
        if not convert_doc:
            msg = ".doc detectado y la conversión automática está desactivada."
            _write_error_md(md_path, folder_name, source_path.name, msg)
            return False, msg

        temp_dir = TEMP_DIR / f"{slugify_ascii(folder_name)}_{os.getpid()}"
        ok_lo, docx_path, msg_lo = convert_to_docx_via_libreoffice(source_path, temp_dir)
        if not ok_lo:
            _write_error_md(md_path, folder_name, source_path.name, msg_lo)
            return False, msg_lo
        ok, md_or_err = docx_to_markdown(docx_path, folder_name)
        if ok:
            md_path.write_text(md_or_err, encoding="utf-8")
            return True, "Word .doc convertido vía LibreOffice → .docx → Markdown."
        _write_error_md(md_path, folder_name, source_path.name, md_or_err)
        return False, md_or_err

    return False, "Extensión no manejada."


def _write_error_md(md_path: Path, folder_name: str, archivo: str, detalle: str) -> None:
    """Escribe un Markdown de error cuando la conversión falla."""
    contenido = (
        f"# {folder_name}\n\n"
        "## Estado de conversion\n\n"
        f"- Archivo detectado: `{archivo}`\n"
        "- Resultado: ERROR\n"
        f"- Detalle: {detalle}\n"
    )
    md_path.write_text(contenido, encoding="utf-8")


# ================================================================
#  EXCEL → SQLITE (sin cambios — pandas es multiplataforma)
# ================================================================

def excel_to_sqlite(excel_path: Path, conn: sqlite3.Connection,
                    table_name: str) -> tuple[bool, int, str]:
    suffix = excel_path.suffix.lower()
    if suffix not in (".xls", ".xlsx"):
        return False, 0, f"Extensión Excel no soportada: {suffix}"
    try:
        engine = "xlrd" if suffix == ".xls" else "openpyxl"
        df = pd.read_excel(excel_path, engine=engine)
    except Exception as exc:
        return False, 0, f"No se pudo leer Excel: {exc}"

    cleaned = [sanitize_column(c) for c in df.columns.tolist()]
    df.columns = make_unique_columns(cleaned)
    conn.execute(f'DROP TABLE IF EXISTS "{table_name}"')
    df.to_sql(table_name, conn, if_exists="replace", index=False)
    return True, int(len(df)), "Tabla creada/reemplazada."


# ================================================================
#  PROCESAMIENTO POR CARPETA
# ================================================================

def process_folder(folder_path: Path, md_dir: Path, conn: sqlite3.Connection,
                   errors: list[str], results: list[FolderResult],
                   convert_doc: bool) -> None:
    folder_name = folder_path.name
    slug        = slugify_ascii(folder_name)
    md_path     = md_dir / f"{slug}.md"
    result      = FolderResult(folder_name=folder_name, slug=slug, md_path=md_path)

    files      = [p for p in folder_path.iterdir() if p.is_file()]
    word_files = [p for p in files if p.suffix.lower() in (".doc", ".docx")]
    pdf_files  = [p for p in files if p.suffix.lower() == ".pdf"]
    excel_files= [p for p in files if p.suffix.lower() in (".xls", ".xlsx")]

    # Seleccionar fuente principal (Word preferido sobre PDF)
    source_file = None
    if len(word_files) == 1:
        source_file = word_files[0]
    elif len(word_files) == 0 and len(pdf_files) == 1:
        source_file = pdf_files[0]

    if source_file is None:
        msg = (
            "Se esperaba 1 Word (.doc/.docx) o 1 PDF. "
            f"Word={len(word_files)}, PDF={len(pdf_files)}."
        )
        result.errors.append(msg)
        if not md_path.exists():
            _write_error_md(md_path, folder_name, "—", msg)
        result.md_ok, result.md_msg = False, msg
    else:
        print(f"  → {folder_name} ({source_file.suffix})", end=" ... ", flush=True)
        t0 = time.time()
        try:
            ok, msg = convert_source_to_md(source_file, md_path, folder_name, convert_doc)
            result.md_ok, result.md_msg = ok, msg
            if ok:
                print(f"✅ ({time.time()-t0:.1f}s)")
            else:
                print(f"❌ {msg[:60]}")
                result.errors.append(msg)
        except Exception as exc:
            msg = f"Fallo Word/PDF→Markdown: {exc}"
            result.md_ok, result.md_msg = False, msg
            result.errors.append(msg)
            _write_error_md(md_path, folder_name, source_file.name, msg)
            print(f"❌ {msg[:60]}")

    # Excel → SQLite
    if len(excel_files) != 1:
        msg = f"Se esperaba 1 Excel y se encontraron {len(excel_files)}."
        result.sql_ok, result.sql_msg = False, msg
        result.errors.append(msg)
    else:
        try:
            ok, rows, msg = excel_to_sqlite(excel_files[0], conn, slug)
            result.sql_ok, result.sql_rows, result.sql_msg = ok, rows, msg
            if not ok:
                result.errors.append(msg)
        except Exception as exc:
            msg = f"Fallo Excel→SQLite: {exc}"
            result.sql_ok, result.sql_msg = False, msg
            result.errors.append(msg)

    if result.errors:
        errors.append(f"[{folder_name}] " + " | ".join(result.errors))
    results.append(result)


# ================================================================
#  REPORTE Y VERIFICACIÓN
# ================================================================

def print_verification(md_dir: Path, conn: sqlite3.Connection,
                        results: list[FolderResult], errors: list[str]) -> None:
    print("\n=== Markdown generados ===")
    for md in sorted(md_dir.glob("*.md")):
        print(f"  {md.name}")

    print("\n=== Tablas SQLite ===")
    for r in [r for r in results if r.sql_ok]:
        count = conn.execute(f'SELECT COUNT(*) FROM "{r.slug}"').fetchone()[0]
        print(f"  {r.slug}: {count} filas")

    total   = len(results)
    md_ok   = sum(1 for r in results if r.md_ok)
    sql_ok  = sum(1 for r in results if r.sql_ok)
    errores = sum(1 for r in results if r.errors)
    print(f"\n=== Resumen ===")
    print(f"  Carpetas procesadas : {total}")
    print(f"  Word/PDF → Markdown : {md_ok} ✅  {total - md_ok} ❌")
    print(f"  Excel   → SQLite    : {sql_ok} ✅  {total - sql_ok} ❌")
    print(f"  Carpetas con errores: {errores}")

    if errors:
        print("\n=== Errores ===")
        for e in errors:
            print(f"  - {e}")


def write_log(log_path: Path, errors: list[str]) -> None:
    log_path.parent.mkdir(parents=True, exist_ok=True)
    if not errors:
        log_path.write_text("Sin errores.\n", encoding="utf-8")
        return
    lines = ["Errores de conversion:\n"] + [f"- {e}\n" for e in errors]
    log_path.write_text("".join(lines), encoding="utf-8")


# ================================================================
#  FUNCIÓN PRINCIPAL
# ================================================================

def run(origen: Path, dest_md: Path, db_path: Path, log_path: Path,
        convert_doc: bool, only_folders: list[str] | None = None) -> int:

    if not origen.exists() or not origen.is_dir():
        print(f"ERROR: No existe carpeta origen: {origen}")
        return 1

    dest_md.mkdir(parents=True, exist_ok=True)
    TEMP_DIR.mkdir(parents=True, exist_ok=True)

    if docx is None:
        print("AVISO: python-docx no instalado; .docx no podrá procesarse.")
    if not _libreoffice_disponible():
        print("AVISO: LibreOffice no instalado; .doc y PDF (fallback) no podrán convertirse.")
        print("       Instale: sudo dnf install libreoffice-headless")

    print("=" * 58)
    print("  SIAA — Convertidor Completo")
    print(f"  Origen : {origen}")
    print(f"  MD     : {dest_md}")
    print(f"  SQLite : {db_path}")
    print("=" * 58)

    errors:   list[str]          = []
    results:  list[FolderResult] = []
    only_set = {n.casefold() for n in only_folders} if only_folders else None

    with sqlite3.connect(db_path) as conn:
        subfolders = sorted(
            [p for p in origen.iterdir() if p.is_dir()],
            key=lambda p: p.name.lower()
        )
        if only_set:
            subfolders = [p for p in subfolders if p.name.casefold() in only_set]

        if not subfolders:
            print(f"\n  No se encontraron subcarpetas en: {origen}")
            return 1

        print(f"\n  Procesando {len(subfolders)} carpeta(s):\n")
        for folder in subfolders:
            process_folder(folder, dest_md, conn, errors, results, convert_doc)
        conn.commit()

    print_verification(dest_md, conn, results, errors)
    write_log(log_path, errors)
    print(f"\n  Log: {log_path}")

    # Recargar índice SIAA automáticamente
    if RECARGAR_AUTO and any(r.md_ok for r in results):
        try:
            import requests
            resp  = requests.get(PROXY_RECARGAR, timeout=15)
            total = resp.json().get("total_docs", "?")
            print(f"  [Proxy] Índice recargado: {total} documentos ✓")
        except Exception:
            print(f"  [Proxy] Recarga manual: curl {PROXY_RECARGAR}")

    return 0


# ================================================================
#  CLI
# ================================================================

def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(
        description="SIAA — Conversor institucional: Word/PDF→Markdown y Excel→SQLite"
    )
    p.add_argument("--origen",   type=Path, default=DEFAULT_ORIGEN,  help="Carpeta raíz con subcarpetas de instructivos.")
    p.add_argument("--dest-md",  type=Path, default=DEFAULT_DEST_MD, help="Destino de archivos Markdown.")
    p.add_argument("--db",       type=Path, default=DEFAULT_DB,      help="Base de datos SQLite.")
    p.add_argument("--log",      type=Path, default=DEFAULT_LOG,     help="Archivo de log de errores.")
    p.add_argument(
        "--convert-doc-to-docx",
        action=argparse.BooleanOptionalAction,
        default=True,
        help="Convierte .doc a .docx con LibreOffice antes de generar Markdown.",
    )
    p.add_argument(
        "--only-folder",
        action="append",
        default=[],
        help="Procesa solo carpetas con este nombre exacto. Repetir para múltiples.",
    )
    return p.parse_args()


if __name__ == "__main__":
    args = parse_args()
    raise SystemExit(
        run(
            args.origen,
            args.dest_md,
            args.db,
            args.log,
            args.convert_doc_to_docx,
            args.only_folder or None,
        )
    )
